"""
format_apply.py  ·  AS9100 컨설팅 자료 일괄 서식 변환 스크립트
──────────────────────────────────────────────────────────────
역할
  1. reference.docx 의 스타일 XML 을 대상 .docx 에 병합
  2. 페이지 여백·용지 크기 적용
  3. 헤더/푸터 중 회사 관련 텍스트 교체
  4. 본문·표 내 회사명·주소·연락처 교체
  5. 결과를 OUTPUT_DIR 에 폴더 구조 그대로 저장
  6. 처리 이력을 LOG_FILE 에 기록

사용법
  C:\\...\\Python313\\python.exe format_apply.py

주의
  - 로고 이미지 교체는 자동화 불가 → 출력 파일에서 수동 교체 필요
  - 아래 설정 섹션의 COMPANY_* 값을 실제 M&C 정보로 먼저 수정하세요
"""

from __future__ import annotations
import copy
import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Mm
from lxml import etree

# ═══════════════════════════════════════════════════
# ★ 설정 섹션 — 실행 전 반드시 확인/수정
# ═══════════════════════════════════════════════════

BASE_DIR       = Path(r"C:\Users\MISTOP\Documents\02_AI디지털전환\AS9100D문서관리")
TARGET_DIR     = BASE_DIR / "AS9100 컨설팅 자료"
REFERENCE_FILE = BASE_DIR / "06_산출물_문서" / "Reference Doc" / "reference.docx"
OUTPUT_DIR     = BASE_DIR / "AS9100 컨설팅 자료_Output"
LOG_FILE       = BASE_DIR / "scripts" / "format_apply_log.txt"

# 교체할 원본 회사 패턴 (왼쪽부터 순서대로 적용)
# 실제 문서에서 확인된 회사명: ㈜영민에프엔에스, ㈜우주항공산업, (주)항공우주산업
REPLACE_PATTERNS: list[tuple[str, str]] = [
    # ㈜영민에프엔에스 (YM 문서 계열) — ㈜ 와 (주) 두 가지 표기 모두 처리
    (r"[㈜]\s*영민에프엔에스",        "M&C Electronics Vina"),
    (r"\(주\)\s*영민에프엔에스",       "M&C Electronics Vina"),
    (r"영민에프엔에스",                "M&C Electronics Vina"),
    # ㈜우주항공산업 (WA 문서 계열)
    (r"[㈜]\s*우주항공산업",           "M&C Electronics Vina"),
    (r"\(주\)\s*우주항공산업",         "M&C Electronics Vina"),
    (r"우주항공산업",                  "M&C Electronics Vina"),
    # (주)항공우주산업 변형
    (r"[㈜]\s*항공우주산업",           "M&C Electronics Vina"),
    (r"\(주\)\s*항공우주산업",         "M&C Electronics Vina"),
    (r"항공우주산업",                  "M&C Electronics Vina"),
    # 주소 · 연락처 — 실제 값으로 교체하세요
    # (r"기존 주소 텍스트", "베트남 법인 실제 주소"),
    # (r"기존 전화번호",    "+84-XX-XXXX-XXXX"),
]

# M&C 회사 기본 정보 (헤더/푸터에 삽입될 수 있음)
COMPANY_KO   = "(주)M&C"
COMPANY_EN   = "M&C Electronics Vina"
COMPANY_ADDR = "[베트남 법인 주소 — 여기에 실제 주소 입력]"  # ← 수정 필요

# ═══════════════════════════════════════════════════
# 유틸리티
# ═══════════════════════════════════════════════════

_compiled_patterns: list[tuple[re.Pattern, str]] = [
    (re.compile(pat, re.IGNORECASE), repl)
    for pat, repl in REPLACE_PATTERNS
]


def _replace_in_text(text: str) -> str:
    for pattern, repl in _compiled_patterns:
        text = pattern.sub(repl, text)
    return text


def _replace_in_paragraph(para) -> bool:
    """단락 내 회사명 교체. 교체 발생 시 True 반환."""
    changed = False
    for run in para.runs:
        new_text = _replace_in_text(run.text)
        if new_text != run.text:
            run.text = new_text
            changed = True
    return changed


def _replace_in_table(table) -> int:
    """표 전체 교체. 교체된 셀 수 반환."""
    count = 0
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if _replace_in_paragraph(para):
                    count += 1
    return count


def _replace_in_header_footer(section) -> int:
    count = 0
    for hf in [section.header, section.footer,
               section.first_page_header, section.first_page_footer,
               section.even_page_header, section.even_page_footer]:
        if hf is None:
            continue
        for para in hf.paragraphs:
            if _replace_in_paragraph(para):
                count += 1
        for table in hf.tables:
            count += _replace_in_table(table)
    return count


# ═══════════════════════════════════════════════════
# 스타일 XML 병합
# ═══════════════════════════════════════════════════

def merge_styles(ref_doc: Document, tgt_doc: Document) -> int:
    """reference.docx 에만 있는 스타일을 대상 문서에 추가. 추가된 수 반환."""
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W  = f"{{{ns}}}"

    # python-docx 1.x: doc.styles.element → CT_Styles lxml element
    ref_root = ref_doc.styles.element
    tgt_root = tgt_doc.styles.element

    existing = {el.get(f"{W}styleId") for el in tgt_root.findall(f"{W}style")}

    added = 0
    for style_el in ref_root.findall(f"{W}style"):
        sid = style_el.get(f"{W}styleId")
        if sid and sid not in existing:
            tgt_root.append(copy.deepcopy(style_el))
            added += 1

    return added


# ═══════════════════════════════════════════════════
# 페이지 설정 적용
# ═══════════════════════════════════════════════════

def apply_page_setup(ref_doc: Document, tgt_doc: Document) -> None:
    """reference.docx 첫 섹션의 여백·용지 크기를 대상 모든 섹션에 적용."""
    ref_sec = ref_doc.sections[0]
    for sec in tgt_doc.sections:
        sec.page_width    = ref_sec.page_width
        sec.page_height   = ref_sec.page_height
        sec.top_margin    = ref_sec.top_margin
        sec.bottom_margin = ref_sec.bottom_margin
        sec.left_margin   = ref_sec.left_margin
        sec.right_margin  = ref_sec.right_margin


# ═══════════════════════════════════════════════════
# 단일 파일 처리
# ═══════════════════════════════════════════════════

def process_file(src: Path, dst: Path, ref_doc: Document) -> dict:
    """하나의 .docx 파일을 처리하고 결과 dict 반환."""
    result = {
        "src": str(src),
        "dst": str(dst),
        "status": "ok",
        "styles_added": 0,
        "replacements": 0,
        "error": "",
    }
    try:
        dst.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(src, dst)            # 원본 복사 후 처리
        doc = Document(dst)

        # 1. 스타일 병합
        result["styles_added"] = merge_styles(ref_doc, doc)

        # 2. 페이지 설정
        apply_page_setup(ref_doc, doc)

        # 3. 본문 교체
        rep = 0
        for para in doc.paragraphs:
            if _replace_in_paragraph(para):
                rep += 1
        for table in doc.tables:
            rep += _replace_in_table(table)

        # 4. 헤더/푸터 교체
        for section in doc.sections:
            rep += _replace_in_header_footer(section)

        result["replacements"] = rep

        doc.save(dst)

    except Exception as exc:
        result["status"] = "error"
        result["error"]  = str(exc)
        # 실패 시 dst 삭제
        if dst.exists():
            dst.unlink()

    return result


# ═══════════════════════════════════════════════════
# 메인
# ═══════════════════════════════════════════════════

def main() -> None:
    # 사전 검증
    if not TARGET_DIR.exists():
        sys.exit(f"[오류] 대상 폴더 없음: {TARGET_DIR}")
    if not REFERENCE_FILE.exists():
        sys.exit(f"[오류] 레퍼런스 파일 없음: {REFERENCE_FILE}")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    print(f"레퍼런스: {REFERENCE_FILE}")
    print(f"대상 폴더: {TARGET_DIR}")
    print(f"출력 폴더: {OUTPUT_DIR}")

    ref_doc = Document(REFERENCE_FILE)

    # 대상 .docx 목록 수집 (임시 파일 제외)
    docx_files = [
        p for p in TARGET_DIR.rglob("*.docx")
        if not p.name.startswith("~$")
    ]
    total = len(docx_files)
    print(f"\n처리 대상: {total}건\n" + "-" * 50)

    log_lines = [
        f"# format_apply.py  실행: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"# 레퍼런스: {REFERENCE_FILE}",
        f"# 대상: {TARGET_DIR}  ({total}건)",
        "상태,스타일추가,교체건수,파일,오류",
    ]

    ok_count = err_count = 0
    for i, src in enumerate(docx_files, 1):
        rel   = src.relative_to(TARGET_DIR)
        dst   = OUTPUT_DIR / rel
        res   = process_file(src, dst, ref_doc)
        icon  = "OK" if res["status"] == "ok" else "NG"
        print(f"[{i:03d}/{total}] {icon} {rel}")
        if res["status"] == "error":
            print(f"       >> {res['error']}")
            err_count += 1
        else:
            ok_count += 1
        log_lines.append(
            f"{res['status']},{res['styles_added']},{res['replacements']},"
            f"\"{rel}\",\"{res['error']}\""
        )

    # 로그 저장
    LOG_FILE.write_text("\n".join(log_lines), encoding="utf-8")

    print("\n" + "=" * 50)
    print(f"완료: {ok_count}건  /  오류: {err_count}건  /  전체: {total}건")
    print(f"출력 위치: {OUTPUT_DIR}")
    print(f"로그 파일: {LOG_FILE}")
    if COMPANY_ADDR.startswith("["):
        print("\n[주의] COMPANY_ADDR 가 플레이스홀더입니다. 스크립트 상단 설정을 수정하세요.")
    print("[주의] 로고 이미지는 자동 교체되지 않습니다. 출력 파일에서 수동 교체하세요.")


if __name__ == "__main__":
    main()
